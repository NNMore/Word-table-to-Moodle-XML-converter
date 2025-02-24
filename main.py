import logging
from docx import Document
import pandas as pd

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def read_document(file_path):
    """Read the Word document and return the first table."""
    try:
        doc = Document(file_path)
        if not doc.tables:
            logging.error("No tables found in the document.")
            return None
        return doc.tables[0]
    except Exception as e:
        logging.error(f"Error reading the document: {e}")
        return None

def extract_text(cell):
    """Extract text from a cell, accounting for empty values."""
    return cell.text.strip() if cell.text.strip() else None

def process_table(table):
    """Process the table and extract questions and answers."""
    data = []
    row_count = len(table.rows)

    # Ensure that the number of rows is a multiple of 5
    if row_count % 5 != 0:
        logging.warning(f"The number of rows ({row_count}) is not a multiple of 5. Some questions may be skipped.")

    for i in range(0, row_count, 5):  # Reading 5 rows (5 rows per question)
        try:
            question_row = table.rows[i]
            question_number = extract_text(question_row.cells[0])
            question = extract_text(question_row.cells[1])

            # If the question number is missing, use the index as the question number
            if not question_number:
                question_number = str(i // 5 + 1)

            answers = []
            for j in range(0, 5):  # Extracting 5 rows with answer options
                try:
                    answer_row = table.rows[i + j]
                    answer_text = extract_text(answer_row.cells[2])
                    is_bold = any(run.bold for run in answer_row.cells[2].paragraphs[0].runs)
                    correct_indicator = '*' if is_bold else ''
                    
                    # Removing indicators before the answer text (letter, dot, comma, and space)
                    if '.' in answer_text:
                        answer_text = answer_text.split('.', 1)[1].strip()
                    elif ',' in answer_text:
                        answer_text = answer_text.split(',', 1)[1].strip()

                    answers.append(f"{correct_indicator}{answer_text}")
                except Exception as e:
                    logging.warning(f"Error processing answer row {i + j}: {e}")
                    continue  # Skip to the next answer if there's an error

            formatted_answers_str = '\n'.join(answers)
            data.append([question_number, question, formatted_answers_str])
        except IndexError:
            # Error handling for incomplete data with question number output
            remaining_rows = len(table.rows) - i
            question_number = extract_text(table.rows[i].cells[0]) if i < len(table.rows) else str(i // 5 + 1)
            logging.warning(f"Not enough rows for question number {question_number}, starting from row {i}. Remaining rows: {remaining_rows}. Skipping.")
        except Exception as e:
            logging.error(f"Error processing question starting at row {i}: {e}")

    return data

def create_xml(data, output_file):
    """Create an XML file for Moodle from the extracted data."""
    try:
        df = pd.DataFrame(data, columns=['Question Number', 'Question Content', 'Answer Options'])
        with open(output_file, 'w', encoding='utf-8') as file:
            file.write('<quiz>\n')
            for index, row in df.iterrows():
                try:
                    file.write('  <question type="multichoice">\n')
                    file.write(f'    <name>\n      <text>Question {row["Question Number"]}</text>\n    </name>\n')
                    file.write(f'    <questiontext format="html">\n      <text><![CDATA[{row["Question Content"]}]]></text>\n    </questiontext>\n')
                    answers = row["Answer Options"].split('\n')
                    for answer in answers:
                        correct = '100' if answer.startswith('*') else '0'
                        answer_text = answer.lstrip('*')
                        file.write(f'    <answer fraction="{correct}">\n')
                        file.write(f'      <text><![CDATA[{answer_text}]]></text>\n')
                        file.write('    </answer>\n')
                    file.write('  </question>\n')
                except Exception as e:
                    logging.warning(f"Error writing question {row['Question Number']} to XML: {e}")
                    continue  # Skip to the next question if there's an error
            file.write('</quiz>')
        logging.info(f"XML file '{output_file}' created successfully.")
    except Exception as e:
        logging.error(f"Error creating XML file: {e}")

def main(file_path, output_file):
    """Main function to read the document, process the table, and create the XML file."""
    table = read_document(file_path)
    if table:
        data = process_table(table)
        create_xml(data, output_file)

if __name__ == "__main__":
    input_file_path = R"path\to\docx\file"  # Update this path
    output_file_path = 'questions.xml'
    main(input_file_path, output_file_path)
